import json
import docx
from docx import Document
from docx.shared import Pt
import os

# Detector recipe files
recipe_files = [
    'all_recipes_threat_detector.json',
    'all_recipes_instance_security.json',
    'all_recipes_config_detector.json',
    'all_recipes_activity_detector.json'
]

# Problem details files prefix
problem_prefix = 'problem.details.'

# Output file
output_file = 'Detector_Recipes_Compliance_Status.docx'

# Risk level order
risk_level_order = {
    'critical': 4,
    'high': 3,
    'medium': 2,
    'low': 1,
    'minor': 0
}

# Create Word document
document = Document()

# Function to set font size and prevent text wrap
def format_cell(cell, text, font_size=6):
    cell.text = text
    run = cell.paragraphs[0].runs[0]
    run.font.size = Pt(font_size)
    # Prevent text wrap by setting word wrap to false
    tcPr = cell._element.get_or_add_tcPr()
    no_wrap = docx.oxml.shared.OxmlElement('w:noWrap')
    tcPr.append(no_wrap)

# Process detector recipe files
for file in recipe_files:
    document.add_heading(f'Detector Recipes for {file}', level=1)

    with open(file, 'r') as recipe_file:
        data = json.load(recipe_file)
        recipes = data['data']['items']

        compliant_recipes = []
        non_compliant_recipes = []

        for recipe in recipes:
            if 'detector-details' in recipe and recipe['detector-details'] is not None:
                status = 'Compliant'
                problem_details = {}

                # Check for problems
                for filename in os.listdir('.'):
                    if filename.startswith(problem_prefix):
                        with open(filename, 'r') as problem_file:
                            problem_data = json.load(problem_file)
                            problem_id = problem_data['data']['detector-rule-id']
                            detector_id = problem_data['data']['detector-id']
                            lifecycle_detail = problem_data['data']['lifecycle-detail']
                            lifecycle_state = problem_data['data']['lifecycle-state']

                            if problem_id == recipe['id'] and detector_id == recipe['detector'] and lifecycle_detail == 'OPEN' and lifecycle_state == 'ACTIVE':
                                status = 'Non-Compliant'
                                problem_details = problem_data['data']

                # Add recipe to corresponding list
                recipe_data = {
                    'Detector ID': recipe['detector'],
                    'Detector Rule ID': recipe['id'],
                    'Labels': recipe['detector-details'].get('labels', []),
                    'Risk-Level': recipe['detector-details']['risk-level'],
                    'Status': status
                }

                if status == 'Compliant':
                    compliant_recipes.append(recipe_data)
                else:
                    non_compliant_recipes.append({'recipe': recipe_data, 'problem_details': problem_details})

        # Sort recipes by risk-level
        compliant_recipes.sort(key=lambda x: risk_level_order.get(x.get('Risk-Level', 'minor') or 'minor', 0), reverse=True)
        non_compliant_recipes.sort(key=lambda x: risk_level_order.get((x['recipe'].get('Risk-Level') or 'minor').lower(), 0), reverse=True)

        # Add compliant recipes table
        document.add_heading('Compliant Recipes', level=2)
        compliant_table = document.add_table(rows=len(compliant_recipes) + 1, cols=5, style='Table Grid')
        hdr_cells = compliant_table.rows[0].cells
        format_cell(hdr_cells[0], 'Detector ID')
        format_cell(hdr_cells[1], 'Detector Rule ID')
        format_cell(hdr_cells[2], 'Labels')
        format_cell(hdr_cells[3], 'Risk-Level')
        format_cell(hdr_cells[4], 'Status')

        for i, recipe in enumerate(compliant_recipes):
            row_cells = compliant_table.rows[i + 1].cells
            risk_level = recipe.get('Risk-Level')
            format_cell(row_cells[0], recipe['Detector ID'])
            format_cell(row_cells[1], recipe['Detector Rule ID'])
            format_cell(row_cells[2], ', '.join(recipe.get('Labels', [])) if isinstance(recipe.get('Labels'), list) else str(recipe.get('Labels', '')))
            format_cell(row_cells[3], 'Unknown' if risk_level is None else risk_level)
            format_cell(row_cells[4], recipe['Status'])

        # Add non-compliant recipes table
        document.add_heading('Non-Compliant Recipes', level=2)
        non_compliant_table = document.add_table(rows=len(non_compliant_recipes) + 1, cols=5, style='Table Grid')
        hdr_cells = non_compliant_table.rows[0].cells
        format_cell(hdr_cells[0], 'Detector ID')
        format_cell(hdr_cells[1], 'Detector Rule ID')
        format_cell(hdr_cells[2], 'Labels')
        format_cell(hdr_cells[3], 'Risk-Level')
        format_cell(hdr_cells[4], 'Status')

        for i, data in enumerate(non_compliant_recipes):
            recipe = data['recipe']
            row_cells = non_compliant_table.rows[i + 1].cells
            risk_level = recipe.get('Risk-Level')
            format_cell(row_cells[0], recipe['Detector ID'])
            format_cell(row_cells[1], recipe['Detector Rule ID'])
            format_cell(row_cells[2], ', '.join(recipe.get('Labels', [])) if isinstance(recipe.get('Labels'), list) else str(recipe.get('Labels', '')))
            format_cell(row_cells[3], 'Unknown' if risk_level is None else risk_level)
            format_cell(row_cells[4], recipe['Status'])

# Save document
document.save(output_file)

print(f"Detector Recipes Status report generated successfully! Saved to {output_file}")