import json
import glob
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement

# Constants
PROBLEM_PREFIX = 'problem.details.*'
OUTPUT_FILE = 'Problem_Details_Report.docx'

def set_font_size(cell, font_size):
    """Set font size for cell text."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(font_size)

def set_no_wrap(cell):
    """Set no-wrap on a cell by modifying the XML directly."""
    tc_pr = cell._element.get_or_add_tcPr()
    no_wrap = OxmlElement('w:noWrap')
    tc_pr.append(no_wrap)

def add_table_for_problem(document, problem_data, serial_number):
    """Adds a problem details table to the document with a serial number."""
    detector_rule_id = problem_data['data'].get('detector-rule-id', 'Unknown')
    document.add_paragraph(f"{serial_number}. Problem details for {detector_rule_id}", style='Heading 2')
    
    # Create a table with two columns for key-value pairs
    table = document.add_table(rows=1, cols=2, style='Table Grid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Key'
    hdr_cells[1].text = 'Value'
    
    # Set font size and no-wrap for header cells
    set_font_size(hdr_cells[0], 6)
    set_font_size(hdr_cells[1], 6)
    set_no_wrap(hdr_cells[0])

    # Populate table with problem details
    for key, value in problem_data['data'].items():
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = json.dumps(value) if isinstance(value, (dict, list)) else str(value)
        
        # Set font size and no-wrap for each cell in the key column
        set_font_size(row_cells[0], 6)
        set_font_size(row_cells[1], 6)
        set_no_wrap(row_cells[0])  # Apply no-wrap to the key column

    # Add a page break after each problem
    document.add_page_break()

def main():
    # Create a new Document
    document = Document()
    
    # Find all JSON files matching the prefix pattern
    problem_files = glob.glob(PROBLEM_PREFIX)
    
    # Loop through files with serial numbering for problem headers
    for i, filename in enumerate(problem_files, start=1):
        try:
            # Load JSON data
            with open(filename, 'r') as problem_file:
                problem_data = json.load(problem_file)
                
                # Add problem details table to the document with serial number
                add_table_for_problem(document, problem_data, i)
                
        except Exception as e:
            print(f"Error processing {filename}: {str(e)}")
    
    # Save document
    document.save(OUTPUT_FILE)
    print(f"Document generated successfully: {OUTPUT_FILE}")

if __name__ == "__main__":
    main()